# Import for working with file-system.
import os
# Import for the GUI
import Tkinter as tk
import tkMessageBox
# Import for working with word documents
from docx import Document
import getFullDocText as fullDoc

def merge():
    finalQuestionsDoc = Document()
    finalQuotesDoc = Document()
    finalStoriesDoc = Document()
    for root, dirs, files in os.walk(os.getcwd()):
        for d in dirs:
            if 'Shiv' in d:
                os.chdir(os.path.abspath(d))
                for root, dirs, files in os.walk(os.getcwd()):
                    for name in files:
                        print('Name: ', name)
                        # print('name', file)
                        if name.endswith('.docx') and 'final' not in name:
                            if 'Questions' in name:
                                fullText = fullDoc.getText(os.path.join(root, name))
                                finalQuestionsDoc.add_heading(name[:-5])
                                finalQuestionsDoc.add_paragraph(name)
                            if 'Quotes' in name:
                                fullText = fullDoc.getText(os.path.join(root, name))
                                finalQuotesDoc.add_heading(name[:-5])
                                finalQuotesDoc.add_paragraph(fullText)
                            if 'Stories' in name:
                                fullText = fullDoc.getText(os.path.join(root, name))
                                finalStoriesDoc.add_heading(name[:-5])
                                finalStoriesDoc.add_paragraph(fullText)
                # current = os.getcwd() + '/' + d;
                print(os.getcwd())
                final_docs_folder = os.getcwd() + '/final_docs/'
                if not os.path.exists(final_docs_folder):
                    # os.chdir(current)
                    os.makedirs('final_docs', mode=0777)
                finalQuestionsDoc.save(final_docs_folder + 'finalQuestions.docx')
                finalQuotesDoc.save(final_docs_folder + 'finalQuotes.docx')
                finalStoriesDoc.save(final_docs_folder + 'finalStories.docx')
#     tkMessageBox.showinfo("Docu-Merger", "Documents merged!")
#     window.destroy()

# window = tk.Tk()
# window.wm_title('Docu-Merger')
# window.geometry('200x100')
# window.configure(background='#146eff')
# button = tk.Button(window, text="Merge Document", command = merge)
# button.pack()
# window.mainloop()

merge()